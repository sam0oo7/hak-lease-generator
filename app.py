import streamlit as st
from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import subprocess, os, tempfile
from datetime import datetime, date


# ─── SIMPLE AUTH ────────────────────────────────────────────────────────────
# hard-coded credentials (don’t use in prod!)
USERNAME = "hak"
PASSWORD = "1234"

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.title("🔒 Please log in")
    user = st.text_input("Username")
    pwd  = st.text_input("Password", type="password")
    if st.button("Log in"):
        if user == USERNAME and pwd == PASSWORD:
            st.session_state.logged_in = True
            st.rerun()
        else:
            st.error("❌ Invalid credentials")
    st.stop()
# ─────────────────────────────────────────────────────────────────────────────


import shutil

# if the 'soffice' command exists in PATH, use that (Cloud/Linux),
# otherwise assume we're on Windows and use your existing path.
if shutil.which("soffice"):
    LIBREOFFICE = "soffice"
else:
    LIBREOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"



from renderer import render_document

# ─── FULL AGREEMENT TEMPLATE ─────────────────────────────────────────────
LEASE_TEMPLATE = """Commercial Lease Agreement

THIS LEASE (this "Lease") dated this {lease_day} day of {lease_month}, {lease_year}

BETWEEN:

{landlord_name} of {landlord_company}
Telephone: {landlord_phone}
Fax: ___________
(the "Landlord")
OF THE FIRST PART
-AND-

{tenant_name} of {tenant_company}
(the "Tenant")
OF THE SECOND PART
[LINE_BREAK]

IN CONSIDERATION OF the Landlord leasing certain premises to the Tenant, the Tenant leasing those premises from the Landlord and the mutual benefits and obligations set forth in this Lease, the receipt and sufficiency of which consideration is hereby acknowledged, the Parties to this Lease (the "Parties") agree as follows:
[LINE_BREAK]

Basic Terms

1. The following basic terms are hereby approved by the Parties and each reference in this Lease to any of the basic terms will be construed to include the provisions set forth below as well as all of the additional terms and conditions of the applicable sections of this Lease where such basic terms are more fully set forth:

a. Landlord: {landlord_company}
b. Address of: {landlord_address}
c. Tenant: {tenant_name}
d. Address of: {tenant_address}
e. Company Number of: {tenant_company_number}
f. Commencement Date of Lease: {lease_commencement_date}
[PAGE_BREAK]
g. Base Rent: £{base_rent} payable per month  
     (Landlord’s Note: {base_rent_note})
h. Permitted Use Address of Premises: {permitted_use_address}
i. Advance Rent: First and last month’s rent
j. Security/Damage Deposit: £{security_deposit}

Let Premises
2. The Landlord agrees to rent to the Tenant the Premises for only the permitted use (the "Permitted Use") of: {permitted_use}

3. No pets or animals are allowed to be kept in or about the Premises or in any common areas in the building containing the Premises. Upon thirty (30) days notice, the Landlord may revoke any consent previously given under this clause.

4. The Landlord reserves the right for itself and for all persons authorized by it, to erect, use and maintain wiring, mains, pipes and conduits and other means of distributing services in and through the Premises, and at all reasonable times to enter upon the Premises for the purpose of installation, maintenance or repair, and such entry will not be an interference with the Tenant’s possession.

Under this Lease
5. The Landlord reserves the right, when necessary, by reason of accident or in order to make repairs, alterations or improvements relating to the Premises or to other portions of the Building to cause temporary obstruction to the Common Areas and Facilities as reasonably necessary and to interrupt or suspend the supply of electricity, water and other services to the Premises until the repairs, alterations or improvements have been completed. There will be no abatement in rent because of such obstruction, interruption or suspension provided that the repairs, alterations or improvements are made as expeditiously as is reasonably possible.

Term
6. The term of the Lease is a periodic tenancy commencing on {lease_commencement_date} and continuing for the period of {lease_duration} ({lease_duration}) years

Security of Tenure
7. The Landlord and the Tenant have agreed that sections 24 to 28 of the Landlord and Tenant Act 1954 do not apply to this Lease.

8. The Landlord has served on the Tenant a notice in the form, or substantially in the form, set out in Schedule 1 to the Regulatory Reform (Business Tenancies) (England and Wales) Order 2003 ("the Order").

9. The requirements specified in Schedule 2 to the Order have been met in that the Tenant has made the appropriate declaration in the form, or substantially in the form, set out in Schedule 2 to the Order.

Break Clause
10. Either the Landlord or the Tenant may give to the other not less than 3 months written notice to end or renew the tenancy just before the end of the given term in this agreement.

Rent
11. Subject to the provisions of this Lease, the Tenant will pay a base rent of £{base_rent} payable per month, for the Premises (the "Base Rent"), without setoff, abatement or deduction. In addition to the Base Rent, the Tenant will pay for any fees or taxes arising from the Tenant’s business.

12. The Tenant will pay the Base Rent on or before the {rent_due_day} of each and every month of the Term to the Landlord.

13. The Tenant will be charged an additional amount of 3% above the base minimum lending rate of the Bank of England on any rent or other money lawfully due which is in arrears for more than fourteen days after the day on which it became due.
    The Tenant will indemnify the Landlord for all rates, taxes, assessments, duties, charges, impositions and outgoings imposed on the Premises, or any owner or occupier of it, during the Term, including all costs reasonably incurred by the Landlord in connection with appealing against the rating evaluation of the Premises.

14. All amounts payable by the Tenant under this Lease are exclusive of any Value Added Tax that may be payable from time to time, and the Tenant will pay to the Landlord the amount of Value Added Tax charged to the Landlord for a Value Added Tax supply made in respect of the Premises that is not recoverable from HM Revenue and Customs.

15. No acceptance by the Landlord of any amount less than the full amount owed will be taken to operate as a waiver by the Landlord for the full amount or in any way to defeat or affect the rights and remedies of the Landlord to pursue the full amount.

Rent Review

16. The Landlord and Tenant will review the Base Rent after each {review_years}-year period, which Base Rent will become effective.

17. The rent review will assume that:
a. the Premises are fit for occupation by a willing tenant;
b. the Premises may be lawfully let for the Permitted Use;
c. the Landlord and Tenant have complied with their respective obligations in this Lease; and

18. If the Landlord and Tenant cannot agree on the Base Rent on the date of review, either the Landlord or Tenant may request the President of the Royal Institution of Chartered Surveyors to appoint an arbitrator to determine the Open Market Rent, in arbitration conducted in accordance with the Arbitration Act 1996, which determination will be binding on the Landlord and Tenant.

19. If the arbitrator is incapable or unwilling to act, the Landlord or Tenant may request the President of the Royal Institution of Chartered Surveyors to appoint a replacement.

20. The results of the rent review will be recorded in a memorandum that will be signed by the Landlord and Tenant.


Use and Occupation

21. The Tenant will carry on business under the name of {business_name}, and will not change such name without the prior written consent of the Landlord, such consent not to be unreasonably withheld. The Tenant will open the whole of the Premises for business to the public fully fixtured, stocked and staffed on the date of commencement of the term and throughout the term, and will continuously occupy and utilize the entire Premises in the active conduct of its business in a reputable manner on such days and during such hours of business as may be determined from time to time by the Landlord.

22. The Tenant covenants that the Tenant will carry on and conduct its business from time to time carried on upon the Premises in such manner as to comply with any statute, including any subordinate legislation, which is in force now or in the future and taking into account any amendment or re-enactment, or any government department, local authority, other public or competent authority or court of competent jurisdiction and of the insurers in relation to the use, occupation and enjoyment of the Building (including in relation to health and safety compliance with the proper practice recommended by all appropriate authorities).

Security Deposit

23. On execution of this Lease, the Tenant will pay the Landlord a security deposit equal to the amount of £{security_deposit} (the “Security Deposit”) to be held by the Landlord without interest. The Landlord will return the Security Deposit to the Tenant at the end of this tenancy, less such deductions as provided in this Lease but no deduction will be made for damage due to reasonable wear and tear.

24. The Tenant may not use the Security Deposit as payment for the Rent.

Default

25. If the Tenant is in default in the payment of any money, whether hereby expressly reserved or deemed as rent, or any part of the rent, and such default continues following any specific due date on which the Tenant is to make such payment, or in the absence of such specific due date, for the 30 days following written notice by the Landlord requiring the Tenant to pay the same then, at the option of the Landlord, this Lease may be terminated upon 30 days notice and the term will then immediately become forfeited and void, and the Landlord may without further notice or any form of legal process immediately reenter the Premises or any part of the Premises and in the name of the whole repossess and enjoy the same as of its former estate anything contained in this Lease or in any statute or law to the contrary notwithstanding.
[PAGE_BREAK]
26. If and whenever:

a. the Tenant’s leasehold interest hereunder, or any goods, chattels or equipment of the Tenant located in the Premises will be taken or seized in execution or attachment, or if any writ of execution will issue against the Tenant or the Tenant will become insolvent or commit an act of bankruptcy or become bankrupt or take the benefit of any legislation that may be in force for bankrupt or insolvent debtor or become involved in voluntary or involuntary winding up, dissolution or liquidation proceedings, or if a receiver will be appointed for the affairs, business, property or revenues of the Tenant; or

b. the Tenant fails to commence, diligently pursue and complete the Tenant’s work to be performed under any agreement to lease pertaining to the Premises or vacate or abandon the Premises, or fail or cease to operate or otherwise cease to conduct business from the Premises, or use or permit or suffer the use of the Premises for any purpose other than as permitted in this clause, or make a bulk sale of its goods and assets which has not been consented to by the Landlord, or move or commence, attempt or threaten to move its goods, chattels and equipment out of the Premises other than in the routine course of its business.

c. then, and in each such case, at the option of the Landlord, this Lease may be terminated without notice and the term will then immediately become forfeited and void, and the Landlord may without notice or any form of legal process immediately reenter the Premises or any part of the Premises and in the name of the whole repossess and enjoy the same as of its former state anything contained in this Lease or in any statute or law to the contrary notwithstanding.

Termination by Landlord
27. In the event that the Landlord has terminated the Lease pursuant to this section, on the expiration of the time fixed in the notice, if any, this Lease and the right, title, and interest of the Tenant under this Lease will terminate in the same manner and with the same force and effect, except as to the Tenant’s liability, as if the date fixed in the notice of cancellation and termination were the end of the Lease.

Distress
28. If and whenever the Tenant is in default in payment of the Rent owed under this Lease, interest on that Rent and any Value Added Tax in respect to that Rent, the Landlord may use the commercial rent arrears recovery (the "CRAR") procedure set out in Schedule 12 (Taking control of goods) of the Tribunals, Courts and Enforcement Act 2007 to have an enforcement agent take control of the Tenant’s goods and sell them to recover that outstanding amount.

Inspections and Landlord’s Right to Enter
29. The Landlord and the Tenant will complete, sign and date a schedule of condition at the beginning and at the end of this tenancy.

30. During the Term and any renewal of this Lease, the Landlord and its agents may enter the Premises to make inspections or repairs at all reasonable times. However, except where the Landlord or its agents consider it is an emergency, the Landlord must have given not less than 24 hours prior written notice to the Tenant.
[PAGE_BREAK]
31. The Tenant acknowledges that the Landlord or its agent will have the right to enter the Premises at all reasonable times to show them to prospective purchasers, encumbrancers, lessees or assignees, and may also during the ninety days preceding the termination of the terms of this Lease, place upon the Premises the usual type of notice to the effect that the Premises are for rent, which notice the Tenant will permit to remain on them.

Utilities and Other Costs
32. The Tenant is responsible for the direct payment of the following utilities and other charges in relation to the Premises: business rates, telephone, Internet and cable.

33. The Tenant may erect, install and maintain a sign of a kind and size in a location, all in accordance with the Landlord’s design criteria for the Building and as first approved in writing by the Landlord. All other signs, as well as the advertising practices of the Tenant, will comply with all applicable rules and regulations of the Landlord. The Tenant will not erect, install or maintain any sign other than in accordance with this section.

Insurance
34. The Tenant is hereby advised and understands that the personal property of the Tenant is not insured by the Landlord for either damage or loss, and the Landlord assumes no liability for any such loss. The Tenant is advised that, if insurance coverage is desired by the Tenant, the Tenant should inquire of Tenant’s insurance agent regarding a Tenant’s Policy of Insurance.

Tenant’s Insurance
35. The Tenant will pay to the Landlord when due the Insurance Charge.

Landlord’s Insurance
36. At the request of the Tenant, the Landlord will produce reasonable evidence of the terms of the Insurance Policies and of payment of the current premium.


Abandonment

37. If at any time during the Term, the Tenant abandons the Premises or any part of the Premises, the Landlord may, at its option, enter the Premises by any means without being liable for any prosecution for such entering, and without becoming liable to the Tenant for damages or for any payment of any kind whatever, and may, at the Landlord’s discretion, as agent for the Tenant, re-let the Premises, or any part of the Premises, for the whole or any part of the then unexpired term, and may receive and collect all rent payable by virtue of such reletting, and, at the Landlord’s option, hold the Tenant liable for any difference between the Rent that would have been payable under this Lease during the balance of the unexpired term, if this Lease had continued in force, and the net rent for such period realized by the Landlord by means of the reletting. If the Landlord’s right of re-entry is exercised following abandonment of the Premises by the Tenant, then the Landlord may consider any personal property belonging to the Tenant and left on the Premises to also have been abandoned, in which case the Landlord may dispose of all such personal property in any manner the Landlord deems proper in compliance with the Torts (Interference with Goods) Act 1977 and is relieved of all liability for doing so.

[PAGE_BREAK]

38. It is agreed between the Landlord and the Tenant that the Landlord will not be liable for any loss, injury, or damage to persons or property resulting from falling plaster, steam, electricity, water, rain, snow or dampness, or from any other cause.

39. It is agreed between the Landlord and the Tenant that the Landlord will not be liable for any loss or damage caused by acts or omissions of other tenants or occupants, their employees or agents, or any persons not the employees or agents of the Landlord, or for any damage caused by the construction of any public or quasi-public works, and in no event will the Landlord be liable for any consequential or indirect damages suffered by the Tenant.

40. It is agreed between the Landlord and the Tenant that the Landlord will not be liable for any loss, injury or damage caused to persons using the Common Areas and Facilities or to vehicles or their contents or any other property on them, or for any damage to property entrusted to its or their employees, or for the loss of any property by theft or otherwise, and all property kept or stored in the Premises will be at the sole risk of the Tenant.

Attorney Fees
41. All costs, expenses and expenditures, including without limitation complete legal costs incurred by the Landlord on a solicitor/client basis as a result of unlawful detainer of the Premises, the recovery of any rent due under the Lease, or any breach by the Tenant of any other condition contained in the Lease, will forthwith upon demand be paid by the Tenant as Additional Rent. All rents, including the Base Rent and Additional Rent, will bear interest at the rate of Twelve percent (12%) per annum from the due date until paid.

Governing Law
42. This Agreement will be construed in accordance with and governed by the laws of England and the Parties submit to the exclusive jurisdiction of the English Courts.

Severability
43. If there is a conflict between any provision of this Lease and the applicable legislation of England (the “Act”), the Act will prevail and such provisions of the Lease will be amended or deleted as necessary in order to comply with the Act. Further, any provisions required by the Act are incorporated into this Lease.

44. If there is a conflict between any provision of this Lease and any form of lease prescribed by the Act, that prescribed form will prevail and such provisions of the Lease will be amended or deleted as necessary in order to comply with that prescribed form. Further, any provisions required by that prescribed form are incorporated into this Lease.

Amendment of Lease
45. Any amendment or modification of this Lease or additional obligation assumed by either party to this Lease in connection with this Lease will only be binding if evidenced in writing signed by each party or an authorized representative of each party.
[PAGE_BREAK]
Assignment and Subletting
46. The Tenant will not assign this Lease in whole or in part, nor sublet all or any part of the Premises, nor grant any licence or part with possession of the Premises or transfer to any other person in whole or in part or any other right or interest under this Lease (except to a parent, subsidiary or affiliate of the Tenant), without the prior written consent of the Landlord in each instance, which consent will not be unreasonably withheld so long as the proposed assignment or sublease complies with the provisions of this Lease.

47. Notwithstanding any assignment or sublease, the Tenant will remain fully liable on this Lease and will not be released from performing any of the terms, covenants and conditions of this Lease.

48. If the Lease is assigned or if the Premises or any part of the Premises are sublet or occupied by anyone other than the Tenant, the Landlord may collect rent directly from the assignee, subtenant or occupant, and apply the net amount collected, or the necessary portion of that amount, to the rent owing under this Lease.

49. The prohibition against assigning or subletting without the consent required by this Lease will be construed to include a prohibition against any assignment or sublease by operation of law.

50. The consent by the Landlord to any assignment or sublease will not constitute a waiver of the necessity of such consent to any subsequent assignment or sublease.

Building Enforcement Action
51. A finding by a competent authority that the Building, or any portion of the Premises, is such that the Tenant must remove from the Building or Premises, will result in termination of this Lease. The Landlord will receive the total of any consequential damages awarded as a result of the finding. All future rent installments to be paid by the Tenant under this Lease will be terminated.

Tenant’s Repairs and Alterations
52. The Tenant covenants with the Landlord to occupy the Premises in a tenant-like manner and not to permit waste. The Tenant will at all times and at its sole expense, subject to the Landlord’s repair obligations, maintain and keep the Premises in good repair, reasonable wear and tear and damage by fire, lightning, tempest or structural hazards excepted. Without limiting the foregoing, the Tenant will keep, repair, replace and maintain all glass, wiring, pipes and mechanical apparatus in, upon or serving the Premises in good and tenantable repair at its sole expense. When it becomes (or, acting reasonably, should have become) aware of any damage, deficiency or defect, the Tenant will notify the Landlord.

53. The Tenant covenants that the Landlord, its servants, agents and workmen may enter and view the state of repair of the Premises and that the Tenant will repair in accordance with any written notice from the Landlord, subject to the Landlord’s repair obligations. If the Tenant refuses or neglects to repair after written demand, the Landlord may, but is not obliged to, undertake such repairs without liability to the Tenant for any loss or damage, and upon completion the Tenant will pay, as Additional Rent, the Landlord’s cost plus fifteen percent (15%) for overhead and supervision.

54. The Tenant will keep in good order and repair the non-structural interior of the Premises and all equipment, fixtures, walls, ceilings, floors, windows, doors, plate glass and skylights. The Tenant will not use or keep any device which might overload the capacity of any floor, wall, utility, electrical or mechanical facility or service.

55. The Tenant will not make or permit alterations, additions or improvements, or install any partitions, trade fixtures, signage, floor coverings, lighting, plumbing, shades, awnings or exterior decorations, without first obtaining the Landlord’s written approval, not to be unreasonably withheld for interior works.

56. The Tenant will not install any special locks, safes or apparatus for air-conditioning, cooling, heating, illuminating, refrigerating or ventilating the Premises, nor add or change locks, without the prior written consent of both the Landlord and the Tenant.

57. When seeking any approval of the Landlord for Tenant repairs as required in this Lease, the Tenant will present to the Landlord plans and specifications of the proposed work which will be subject to the prior approval of the Landlord, not to be unreasonably withheld or delayed.

58. The Tenant will be responsible at its own expense to replace all electric light bulbs, tubes, ballasts or fixtures serving the Premises.

Landlord’s Repairs
59. The Landlord covenants and agrees to effect at its expense repairs of a structural nature to the structural elements of the roof, foundation and outside walls of the Building, whether occasioned or necessitated by faulty workmanship, materials, improper installation, construction defects or settling, or otherwise, unless such repair is necessitated by the negligence of the Tenant, its servants, agents, employees or invitees, in which event the cost of such repairs will be paid by the Tenant together with an administration fee of fifteen percent (15%) for the Landlord’s overhead and supervision.

Care and Use of Premises
60. The Tenant will promptly notify the Landlord of any damage, or of any situation that may significantly interfere with the normal use of the Premises.

61. The Tenant will not make (or allow to be made) any noise or nuisance which, in the reasonable opinion of the Landlord, disturbs the comfort or convenience of other tenants.

62. The Tenant will dispose of its rubbish in a timely, tidy, proper and sanitary manner.

63. The Tenant will not engage in any illegal trade or activity on or about the Premises.

64. The Landlord and Tenant will comply with standards of health, sanitation, fire, housing and safety as required by law.

65. The hallways, passages and stairs of the building in which the Premises are situated will be used for no purpose other than going to and from the Premises and the Tenant will not in any way encumber those areas with boxes, furniture or other material or place or leave rubbish in those areas and other areas used in common with any other tenant.
[PAGE_BREAK]
Surrender of Premises
66. The Tenant covenants to surrender the Premises, at the expiration of the tenancy created in this Lease, in the same condition as the Premises were in upon delivery of possession under this Lease, reasonable wear and tear, damage by fire or the elements, and unavoidable casualty excepted, and agrees to surrender all keys for the Premises to the Landlord at the place then fixed for payment of rent and will inform the Landlord of all combinations to locks, safes and vaults, if any. All alterations, additions and improvements constructed or installed in the Premises and attached in any manner to the floor, walls or ceiling, including any leasehold improvements, equipment, floor covering or fixtures (including trade fixtures), will remain upon and be surrendered with the Premises and will become the absolute property of the Landlord except to the extent that the Landlord requires removal of such items. If the Tenant abandons the Premises or if this Lease is terminated before the proper expiration of the term due to a default on the part of the Tenant then, in such event, as of the moment of default of the Tenant all trade fixtures and furnishings of the Tenant (whether or not attached in any manner to the Premises) will, except to the extent the Landlord requires the removal of such items, become and be deemed to be the property of the Landlord without indemnity to the Tenant and as liquidated damages in respect of such default but without prejudice to any other right or remedy of the Landlord. Notwithstanding that any trade fixtures, furnishings, alterations, additions, improvements or fixtures are or may become the property of the Landlord, the Tenant will immediately remove all or part of the same and will make good any damage caused to the Premises resulting from the installation or removal of such fixtures, all at the Tenant’s expense, should the Landlord so require by notice to the Tenant. If the Tenant, after receipt of such notice from the Landlord, fails to promptly remove any trade fixtures, furnishings, alterations, improvements and fixtures in accordance with such notice, the Landlord may enter into the Premises and remove from the Premises all or part of such trade fixtures, furnishings, alterations, additions, improvements and fixtures without any liability and at the expense of the Tenant, which expense will immediately be paid by the Tenant to the Landlord. The Tenant’s obligation to observe or perform the covenants contained in this Lease will survive the expiration or other termination of the Term.

Hazardous Materials
67. The Tenant will not keep or have on the Premises any article or thing of a dangerous, flammable, or explosive character that might unreasonably increase the danger of fire on the Premises or that might be considered hazardous by any responsible insurance company.

Rules and Regulations
68. The Tenant will obey all rules and regulations posted by the Landlord regarding the use and care of the Building, car park and other common facilities that are provided for the use of the Tenant in and around the Building on the Premises.

69. The Tenant will obey all rules and regulations posted by the Landlord regarding the use and care of the Building, car park and other common facilities that are provided for the use of the Tenant in and around the Building on the Premises.
[PAGE_BREAK]
Address for Notice
70. For any matter relating to this tenancy, whether during or after this tenancy has been terminated:
    
    a. the address for service of the Tenant is the Premises during this tenancy, and {tenant_notice_address} after this tenancy is terminated. The phone number of the Tenant is {tenant_phone}; and
    
    b. the address for service of the Landlord is {landlord_notice_address} both during this tenancy and after it is terminated. The phone number of the Landlord is {landlord_phone}.

The Landlord or the Tenant may, on written notice to each other, change their respective addresses for notice under this Lease.

Limited Liability Beyond Insurance Coverage
71. Notwithstanding anything contained in this Lease to the contrary, for issues relating to this Lease, presuming the Landlord obtains its required insurance, the Landlord will not be liable for loss of Tenant business income, Tenant moving expenses, and consequential, incidental, punitive and indirect damages which are not covered by the Landlord’s insurance.

Remedies Cumulative
72. No reference to or exercise of any specific right or remedy by the Landlord will prejudice or preclude the Landlord from any other remedy whether allowed at law or in equity or expressly provided for in this Lease. No such remedy will be exclusive or dependent upon any other such remedy, but the Landlord may from time to time exercise any one or more of such remedies independently or in combination.

Landlord May Perform
73. If the Tenant fails to observe, perform or keep any of the provisions of this Lease to be observed, performed or kept by it and such failure is not rectified within the time limits specified in this Lease, the Landlord may, but will not be obliged to, at its discretion and without prejudice, rectify the default of the Tenant. The Landlord will have the right to enter the Premises for the purpose of correcting or remedying any default of the Tenant and to remain until the default has been corrected or remedied. However, any expenditure by the Landlord incurred in any correction of a default of the Tenant will not be deemed to waive or release the Tenant’s default or the Landlord’s right to take any action as may be otherwise permissible under this Lease in the case of any default.
[PAGE_BREAK]
General Provisions

74. The Tenant authorizes the Landlord to make inquiries to any agency related to the Tenant’s compliance with any laws, regulations, or other rules related to the Tenant or the Tenant’s use of the Premises. The Tenant will provide to the Landlord any written authorization that the Landlord may reasonably require to facilitate these inquiries.

75. This Lease will extend to and be binding upon and inure to the benefit of the respective heirs, executors, administrators, successors and assigns, as the case may be, of each party to this Lease. All covenants are to be construed as conditions of this Lease.

76. All sums payable by the Tenant to the Landlord pursuant to any provision of this Lease will be deemed to be Additional Rent and will be recoverable by the Landlord as rental arrears. Where there is more than one Tenant executing this Lease, all Tenants are jointly and severally liable for each other’s acts, omissions and liabilities pursuant to this Lease.

[LINE_BREAK]

1. The provisions of Section 196 of the Law of Property Act 1925, as amended by the Recorded Delivery Services Act 1962, will apply to the giving and service of all notices and documents under or in connection with this Lease.

2. All schedules to this Lease are incorporated into and form an integral part of this Lease.

3. Headings are inserted for the convenience of the Parties only and are not to be considered when interpreting this Lease. Words in the singular mean and include the plural and vice versa. Words in the masculine mean and include the feminine and vice versa.

4. This Lease may be executed in counterparts. Facsimile signatures are binding and are considered to be original signatures.

5. Time is of the essence in this Lease.

6. This Lease will constitute the entire agreement between the Landlord and the Tenant. Any prior understanding or representation of any kind preceding the date of this Lease will not be binding on either party to this Lease except to the extent incorporated in this Lease. No warranties of the Landlord not expressed in this Lease are to be implied.

7. The Parties do not intend for any term of this Lease to be enforceable by a person that is not party to this Lease pursuant to the Contracts (Rights of Third Parties) Act 1999.

8. Nothing contained in this Lease is intended by the Parties to create a relationship of principal and agent, partnership, nor joint venture. The Parties intend only to create a relationship of landlord and tenant.

[PAGE_BREAK]
IN WITNESS WHEREOF the Parties to this Lease have duly affixed their signatures under hand and seal, or by a duly authorised officer under seal, on this {signature_day} day of {signature_month}, {signature_year}
[LINE_BREAK]
[LINE_BREAK]
[SIG_TABLE]

"""

# ─────────────────────────────────────────────────────────────────────────

def format_date(dt: date):
    return {
        "day":   f"{dt.day:02d}",
        "month": dt.strftime("%B"),
        "year":  str(dt.year),
        "full":  dt.strftime("%d/%m/%Y"),
    }

st.set_page_config(page_title="HAK PROPERTIES LEASE GENERATOR")
st.header("📄 HAK PROPERTIES LEASE GENERATOR")

# Header & footer inputs
header_text = st.text_input("Header text", "HAK PROPERTIES LONDON LIMITED")
footer_text = st.text_input("Footer text", "Confidential – For intended recipient only")

with st.form("lease_form"):

    with st.expander("Important Dates"):
        lease_start_str  = st.text_input(
            "Lease Start Date (DD/MM/YYYY)", 
            datetime.today().strftime("%d/%m/%Y")
        )
        # parse it back if you need a date object:
        lease_start = datetime.strptime(lease_start_str, "%d/%m/%Y").date()

        commencement_str = st.text_input(
            "Commencement Date (DD/MM/YYYY)",
            lease_start_str
        )
        commencement = datetime.strptime(
            commencement_str, "%d/%m/%Y"
        ).date()

        signature_str = st.text_input(
            "Signature Date (DD/MM/YYYY)",
            datetime.today().strftime("%d/%m/%Y")
        )
        signature = datetime.strptime(
            signature_str, "%d/%m/%Y"
        ).date()


    with st.expander("🏠 Landlord Details"):
        landlord_name           = st.text_input("Landlord Name", "Samad A Kaka")
        landlord_company        = st.text_input("Landlord Company", "HAK PROPERTIES LONDON LIMITED")
        landlord_phone          = st.text_input("Landlord Phone", "07951209900")
        landlord_address        = st.text_input("Landlord Address", "19, St Michaels Ave, Wembley HA9 6SJ")
        landlord_notice_address = st.text_input("Landlord Notice Address", landlord_address)

    with st.expander("👥 Tenant Details"):
        tenant_name           = st.text_input("Tenant Name", "")
        tenant_company        = st.text_input("Tenant Company", "")
        tenant_address        = st.text_input("Tenant Address", "")
        tenant_company_number = st.text_input("Tenant Company No.", "")
        tenant_signatory      = st.text_input("Tenant Signatory", "")
        tenant_phone          = st.text_input("Tenant Phone", "")
        tenant_notice_address = st.text_input("Tenant Notice Address", "")

    with st.expander("💷 Financials & Use"):
        base_rent             = st.number_input("Base Rent (£/mo)", 0, 10_000, 1_000, step=50)
        base_rent_note        = st.text_input("Landlord’s Note (on Base Rent)", "Additional notes if any")
        permitted_use_address = st.text_input("Permitted Use Address", "")
        permitted_use         = st.text_input("Permitted Use", "")
        security_deposit      = st.number_input("Security Deposit (£)", 0, 10_000, 0, step=50)
        rent_due_day          = st.text_input("Rent Due Day", "1st")
        review_years          = st.number_input("Review Interval (yrs)", 1, 10, 4)
        lease_duration        = st.number_input("Lease Duration (yrs)", 1, 30, 4)

    business_name    = st.text_input("Business Name", tenant_company)
    submitted        = st.form_submit_button("Generate .docx")

if submitted:
    # Date validation
    if commencement < lease_start:
        st.error("🔴 Commencement date cannot be before lease start.")
        st.stop()
    if signature < commencement:
        st.error("🔴 Signature date cannot be before commencement.")
        st.stop()

    # Build context dict
    ls = format_date(lease_start)
    cm = format_date(commencement)
    sg = format_date(signature)
    ctx = {
        # Dates
        "lease_day": ls["day"], "lease_month": ls["month"], "lease_year": ls["year"],
        "lease_commencement_date": cm["full"],
        "signature_day": sg["day"], "signature_month": sg["month"], "signature_year": sg["year"],
        # Landlord
        "landlord_name": landlord_name, "landlord_company": landlord_company,
        "landlord_phone": landlord_phone, "landlord_address": landlord_address,
        "landlord_notice_address": landlord_notice_address,
        # Tenant
        "tenant_name": tenant_name, "tenant_company": tenant_company,
        "tenant_address": tenant_address, "tenant_company_number": tenant_company_number,
        "tenant_signatory": tenant_signatory, "tenant_phone": tenant_phone,
        "tenant_notice_address": tenant_notice_address,
        # Financials
        "base_rent": f"{base_rent:,}", 
        "base_rent_note": base_rent_note,
        "permitted_use": permitted_use,
        "permitted_use_address": permitted_use_address,
        "security_deposit": f"{security_deposit:,}", "rent_due_day": rent_due_day,
        "review_years": str(review_years), "lease_duration": str(lease_duration),
        # Business
        "business_name": business_name,
    }

    # Fill template & create doc
    body = LEASE_TEMPLATE.format(**ctx)
    doc  = Document()

    # 1" margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after  = Pt(6)

    # Header
    hdr = doc.sections[0].header
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    p.text = header_text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer
    ftr = doc.sections[0].footer
    p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    p.text = footer_text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Prepare placeholder lookup
    # new: underline absolutely everything the user typed into the form
    fill_ins = set(ctx.values())

    alternation = "|".join(sorted((re.escape(s) for s in fill_ins), key=len, reverse=True))
    user_pattern = re.compile(rf"(?<!\w)({alternation})(?!\w)")


    # Delegate all rendering & formatting
    render_document(doc, body, fill_ins, user_pattern, ctx)


    # Download button
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        "📥 Download Lease Agreement (.docx)",
        buf,
        "Lease_Agreement.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    

    # ─── PDF Download via LibreOffice with spinner ───────────────────────────────────────
    import subprocess
    from subprocess import CalledProcessError
    import tempfile
    import os

    # 1) dump the docx to a temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_dir, tmp_name = os.path.split(tmp.name)
        pdf_name = os.path.splitext(tmp_name)[0] + ".pdf"
        pdf_path = os.path.join(tmp_dir, pdf_name)

    # 2) headless convert with soffice (LibreOffice), wrapped in a spinner
    with st.spinner("Converting to PDF… please wait"):
        try:
            subprocess.run(
                [
                    LIBREOFFICE,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmp_dir,
                    tmp.name
                ],
                check=True,
                capture_output=True,
                text=True
            )
        except FileNotFoundError:
            st.error("❌ Couldn’t find LibreOffice—make sure `LIBREOFFICE` points to soffice.exe.")
            raise
        except CalledProcessError as e:
            st.error(f"❌ PDF conversion failed:\n{e.stderr}")
            raise

    # 3) read PDF bytes and offer download (only if file exists)
    if os.path.exists(pdf_path):
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        st.download_button(
            "📥 Download Lease Agreement (.pdf)",
            pdf_bytes,
            "Lease_Agreement.pdf",
            "application/pdf"
        )
    else:
        st.error("❌ PDF file not found after conversion.")
