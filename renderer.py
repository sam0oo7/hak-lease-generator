import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# All your section titles
PAGE_TITLES = {
    "Basic Terms","Let Premises","Under this Lease","Term","Security of Tenure","Break Clause",
    "Rent","Use and Occupation","Security Deposit","Default","Distress",
    "Rent Review","Abandonment","Rules and Regulations","Address for Notice",
    "Utilities and Other Costs","Insurance","Tenant’s Insurance","Landlord’s Insurance",
    "Attorney Fees","Governing Law","Severability","Amendment of Lease",
    "Assignment and Subletting","Building Enforcement Action","Tenant’s Repairs and Alterations",
    "Landlord’s Repairs","Care and Use of Premises","Surrender of Premises",
    "Hazardous Materials","General Provisions", "Termination by Landlord", "Inspections and Landlord’s Right to Enter", "Limited Liability Beyond Insurance Coverage",
    "Remedies Cumulative", "Landlord May Perform" 

}


# constants at top of module
PAD = " "
BLANK_LEN = 20
def render_document(doc, body, fill_ins, user_pattern, ctx):
    """
    Renders the lease body into a Word document with correct formatting.
    :param doc: The Document object to add content to.
    :param body: The full lease text with placeholders filled.
    :param fill_ins: A set of strings that should be rendered as fillable blanks.
    :param user_pattern: A compiled regex matching any of the fill_ins values.
    """
    for raw in body.split("\n"):
        line = raw.rstrip()

        # ── Left-align closing clause with fill-ins ─────────────────────
        if line.strip().startswith("IN WITNESS WHEREOF"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for part in user_pattern.split(line):
                if part in fill_ins:

                    # inside your for-part loop:
                    if part.strip():
                        # there is user text → pad it and underline
                        # preserve the one-space indent at the start of the line
                        text = " " + (f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN)

                    else:
                        # empty fill-in → underline a run of spaces
                        text = PAD * BLANK_LEN

                    run = p.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(12)
                    run.underline = True

                else:
                    p.add_run(part)
            continue


        # ─── Signature 2-col table ─────────────────────────────
        if line.strip() == "[SIG_TABLE]":
            # Build a 3×2 table
            tbl = doc.add_table(rows=3, cols=2)
            tbl.autofit = False
            tbl.columns[0].width = Inches(3.25)
            tbl.columns[1].width = Inches(3.25)

            # strip borders
            for cell in tbl._cells:
                for b in cell._element.tcPr.xpath("./w:tcBorders/*"):
                    b.set(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                        "nil"
                    )

            # ── Row 0: landlord’s witness vs landlord, *one line down* ──
            left0 = tbl.cell(0, 0)
            left0.text = "\n\n\n" + (
                "_________________________\n"
                "(Witness Name)\n\n"
                "_________________________\n"
                "(Address)\n\n\n"
                "_________________________\n"
                "(Signature)"
            )
            for p in left0.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            right0 = tbl.cell(0, 1)
            right0.text = "\n\n\n\n\n" + (  # **one line down** for landlord block
                "_________________________\n"
                f"{ctx['landlord_name']}\n"
                f"{ctx['landlord_company']}"
            )
            for p in right0.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # ── Row 1: blank spacing ──
            spacer = tbl.cell(1, 0).merge(tbl.cell(1, 1))
            spacer.text = "\n\n\n"  # two blank lines between landlord and tenant

            # ── Row 2: tenant’s witness vs tenant, *one line down* for the Per: line ──
            tbl.cell(2, 0).text = (
                "_________________________\n"
                "(Witness Name)\n\n"
                "_________________________\n"
                "(Address)\n\n\n"
                "_________________________\n"
                "(Signature)"
            )
            for p in tbl.cell(2, 0).paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            right2 = tbl.cell(2, 1)
            right2.text = "\n\n\n" + (
                "Signed for and on behalf of:\n"
                f"{ctx['tenant_name']}\n\n"
                "\nPer: ________________ (SEAL)"  # **one line down** before Per:
            )
            for p in right2.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            continue           
        # 1) Page break marker
        if line.strip() == "[PAGE_BREAK]":
            doc.add_page_break()
            continue
        # custom single-line break (no extra paragraph spacing)
        if line.strip() == "[LINE_BREAK]":
            # add a soft break to the last paragraph
            last = doc.paragraphs[-1]
            last.add_run().add_break()
            continue
        
        # 2) Skip empty lines
        if not line.strip():
            continue

       

        # 3) Document title
        if line.strip() == "Commercial Lease Agreement":
            p = doc.add_paragraph(line, style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # 4) Bold leading 'THIS LEASE'
        if line.startswith("THIS LEASE"):
            p = doc.add_paragraph()
            p.add_run("THIS LEASE").bold = True
            rest = line[len("THIS LEASE"):]
            for part in user_pattern.split(rest):
                if part in fill_ins:

                    # inside your for-part loop:
                    if part.strip():
                        # there is user text → pad it and underline
                        # preserve the one-space indent at the start of the line
                        text = " " + (f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN)

                    else:
                        # empty fill-in → underline a run of spaces
                        text = PAD * BLANK_LEN

                    run = p.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(12)
                    run.underline = True

                else:
                    p.add_run(part)
            continue

        # 5) BETWEEN: heading spacing
        if line.strip() == "BETWEEN":
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(12)
            continue

        # 5.1) center the “-AND-” separator
        if line.strip() == "-AND-":
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # 6) Telephone + Fax on one centred line
        if line.startswith("Telephone:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ─── Telephone ─────────────────────────────────
            label, _, val = line.partition(":")
            tel_lab = p.add_run(label + ": ")
            tel_lab.bold = True

            tel_text = f"{PAD}{val.strip()}{PAD}" if val.strip() else PAD * BLANK_LEN
            run_tel = p.add_run(tel_text)
            run_tel.font.name = "Courier New"
            run_tel.font.size = Pt(12)
            run_tel.underline = True

             # spacer between telephone and fax
            p.add_run("   ")

            # Fax label + value
            fax_lab = p.add_run("Fax: ")
            fax_lab.bold = True

            fax_val = ctx.get("fax_number", "").strip()
            if fax_val and fax_val != "—":
                # user provided a fax → render it underlined normally
                fax_text = f"{PAD}{fax_val}{PAD}"
                run_fax = p.add_run(fax_text)
                run_fax.font.name = "Courier New"
                run_fax.font.size = Pt(12)
                run_fax.underline = True
            else:
                # blank: build a fixed‐width field of underscores with the dash in the middle
                total_chars = 12
                left  = (total_chars - 1) // 2   # e.g. 5
                right = total_chars - left - 1   # e.g. 6
                # build "__..._-__..."
                blank_field = "_" * left + "-" + "_" * right
                run_fax = p.add_run(blank_field)
                run_fax.font.name = "Courier New"
                run_fax.font.size = Pt(12)
                run_fax.underline = True
                # no run_fax.underline needed, underscores are visible

            # spacer if you still need it
            p.add_run("   ")
            continue
                








        # 7) RIGHT-align uppercase PART text
        if line.strip().endswith("PART") and line.strip().isupper():
            # write the PART line
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # now add two blank lines
            doc.add_paragraph()
            doc.add_paragraph()

            continue


        # 8) Section headings
        if line.strip().rstrip(":") in PAGE_TITLES:
            p = doc.add_paragraph(line.strip(), style="Heading 2")
            r = p.runs[0]
            r.bold = True
            r.font.color.rgb = RGBColor(0, 112, 192)
            continue


        # 8.1) Main numbered clauses with hanging indent
        # strip off any leading spaces so " 10. ..." still matches
        stripped = line.lstrip()
        if re.match(r"^\d+\.\s", stripped):
            # split number from the rest, then trim any extra padding
            number, rest = stripped.split(" ", 1)
            rest = rest.lstrip()

            # start the paragraph with a hanging indent
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)

            # write the clause number
            run = p.add_run(f"{number}  ")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

            # write the clause text, underlining any fill-ins
            for part in user_pattern.split(rest):
                if part in fill_ins:
                    text = f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN
                    r = p.add_run(text)
                    r.font.name = "Courier New"
                    r.font.size = Pt(12)
                    r.underline = True
                else:
                    p.add_run(part)
            continue




        # 8.2) Lettered sub-items: indent under clause
        if re.match(r"^[a-z]\.\s", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            for part in user_pattern.split(line):
                if part in fill_ins:
                    # apply PAD/BLANK_LEN logic
                    if part.strip():
                        text = " " + (f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN)
                    else:
                        text = PAD * BLANK_LEN
                    run = p.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(12)
                    run.underline = True
                else:
                    p.add_run(part)
            continue

        # 9) Centered lines containing fill items
        if " of " in line and any(term in line for term in fill_ins):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for part in user_pattern.split(line):
                if part in fill_ins:

                    # inside your for-part loop:
                    if part.strip():
                        # there is user text → pad it and underline
                        text = " " + (f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN)
                    else:
                        # empty fill-in → underline a run of spaces
                        text = PAD * BLANK_LEN

                    run = p.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(12)
                    run.underline = True

                else:
                    p.add_run(part)
            continue
             # Center the role labels
        if line.strip().startswith("(the"):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

            # (NEW) — totally left-aligned “consideration” paragraph
        if line.startswith("IN CONSIDERATION"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent      = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # then copy in your normal fill-in logic:
            for part in user_pattern.split(line):
                if part in fill_ins:

                    # inside your for-part loop:
                    if part.strip():
                        # there is user text → pad it and underline
                        text = " " + (f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN)
                    else:
                        # empty fill-in → underline a run of spaces
                        text = PAD * BLANK_LEN

                    run = p.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(12)
                    run.underline = True

                else:
                    p.add_run(part)
            continue

        # 10) Fallback: justified paragraph
        p = doc.add_paragraph()
        # exactly match your numbered‐clause indent…
        p.paragraph_format.left_indent      = Inches(0.20)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for part in user_pattern.split(line):
            if part in fill_ins:

                # inside your for-part loop:
                if part.strip():
                    # there is user text → pad it and underline
                    text = " " + (f"{PAD}{part}{PAD}" if part.strip() else PAD * BLANK_LEN)
                else:
                    # empty fill-in → underline a run of spaces
                    text = PAD * BLANK_LEN

                run = p.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(12)
                run.underline = True

            else:
                p.add_run(part)

        

